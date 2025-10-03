from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from io import BytesIO

def set_resume_style(doc):
    """Set global font and paragraph style"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'   # Use 'Cambria' if you prefer more formal
    font.size = Pt(10.5)
    para_format = style.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(2)
    para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

def add_section_header(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    return para

def create_resume_docx(profile, jd_analysis, tailored_summary, match_details, selected_projects, optimized_experiences):
    doc = Document()
    set_margins(doc)
    set_resume_style(doc)

    personal = profile.get('personal', {})

    # Header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(personal.get('name', 'YOUR NAME').upper())
    name_run.bold = True
    name_run.font.size = Pt(16)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = [personal.get(k) for k in ['email', 'phone', 'location', 'linkedin'] if personal.get(k)]
    contact_run = contact_para.add_run(' | '.join(contact_parts))
    contact_run.font.size = Pt(9.5)

    # Professional Summary
    add_section_header(doc, "Professional Summary")
    summary_para = doc.add_paragraph(tailored_summary)
    summary_para.runs[0].font.size = Pt(10.5)

    # Skills
    add_section_header(doc, "Technical Skills")
    all_skills = profile.get('skills', {}).get('technical', [])
    matched_skills = [s for s in all_skills if s.lower() in [m.lower() for m in match_details.get('matched_skills', [])]]
    ordered_skills = matched_skills + [s for s in all_skills if s not in matched_skills]
    skills_para = doc.add_paragraph(', '.join(ordered_skills[:18]))
    skills_para.runs[0].font.size = Pt(10.5)

    # Experience
    if optimized_experiences:
        add_section_header(doc, "Professional Experience")
        for exp in optimized_experiences[:2]:
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(f"{exp.get('role', '')} | {exp.get('company', '')}")
            role_run.bold = True
            role_run.font.size = Pt(11)

            duration_para = doc.add_paragraph(exp.get('duration', ''))
            duration_para.runs[0].italic = True
            duration_para.runs[0].font.size = Pt(9.5)

            for resp in exp.get('responsibilities', [])[:3]:
                bullet = doc.add_paragraph(resp, style='List Bullet')
                bullet.paragraph_format.left_indent = Inches(0.25)
                bullet.runs[0].font.size = Pt(10)

    # Projects
    if selected_projects:
        add_section_header(doc, "Projects")
        for proj in selected_projects[:2]:
            proj_run = doc.add_paragraph().add_run(proj.get('name', 'Project'))
            proj_run.bold = True
            proj_run.font.size = Pt(11)
            desc = proj.get('description', '')
            if len(desc) > 150:
                desc = desc[:150] + '...'
            doc.add_paragraph(desc).runs[0].font.size = Pt(10)
            if proj.get('tech_stack'):
                doc.add_paragraph(f"Tech: {', '.join(proj['tech_stack'][:6])}").runs[0].italic = True

    # Education
    education = profile.get('education', [])
    if education:
        add_section_header(doc, "Education")
        for edu in education[:2]:
            degree_run = doc.add_paragraph().add_run(edu.get('degree', 'Degree'))
            degree_run.bold = True
            degree_run.font.size = Pt(11)
            inst_line = f"{edu.get('institution', '')}"
            if edu.get('year'):
                inst_line += f" | {edu.get('year')}"
            if edu.get('gpa'):
                inst_line += f" | GPA: {edu.get('gpa')}"
            doc.add_paragraph(inst_line).runs[0].font.size = Pt(10)

    # Certifications
    certifications = profile.get('certifications', [])
    if certifications:
        add_section_header(doc, "Certifications")
        certs = [f"{c.get('name','')} ({c.get('issuer','')})" for c in certifications[:3]]
        doc.add_paragraph(' | '.join(certs)).runs[0].font.size = Pt(10)

    # Save DOCX to memory
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
